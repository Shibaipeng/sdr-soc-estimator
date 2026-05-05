"""
按口腔癌论文排版格式生成电池SOC研究论文Word文档
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os, io, copy

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '电池SOC论文_v9_modified.docx')

# ============================================================
# OMML 公式引擎 (Word原生公式, 非图片)
# ============================================================
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
M = f'{{{MATH_NS}}}'

def _m(tag):
    return f'{M}{tag}'

def _mr(text):
    """<m:r> 含 <m:t>"""
    if not isinstance(text, str):
        text = str(text)
    elem = etree.Element(_m('r'))
    t = etree.SubElement(elem, _m('t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return elem

def _mSub(base_text, sub_text):
    """下标: base_sub"""
    elem = etree.Element(_m('sSub'))
    e = etree.SubElement(elem, _m('e'))
    if isinstance(base_text, str):
        e.append(_mr(base_text))
    else:
        e.append(base_text)
    s = etree.SubElement(elem, _m('sub'))
    if isinstance(sub_text, str):
        s.append(_mr(sub_text))
    else:
        s.append(sub_text)
    return elem

def _mFunc(fname, *arg_elems):
    """函数: exp(arg1, arg2, ...)"""
    elem = etree.Element(_m('func'))
    fn = etree.SubElement(elem, _m('fName'))
    fn.append(_mr(fname))
    fe = etree.SubElement(elem, _m('e'))
    for a in arg_elems:
        if isinstance(a, str):
            fe.append(_mr(a))
        else:
            fe.append(a)
    return elem

def _mFrac(num_elem, den_elem):
    """分数"""
    elem = etree.Element(_m('f'))
    n = etree.SubElement(elem, _m('num'))
    n.append(num_elem if not isinstance(num_elem, str) else _mr(num_elem))
    d = etree.SubElement(elem, _m('den'))
    d.append(den_elem if not isinstance(den_elem, str) else _mr(den_elem))
    return elem

def _mParen(*inner_elems):
    """括号"""
    elem = etree.Element(_m('d'))
    etree.SubElement(elem, _m('dPr'))
    fe = etree.SubElement(elem, _m('e'))
    for a in inner_elems:
        if isinstance(a, str):
            fe.append(_mr(a))
        else:
            fe.append(a)
    return elem

def _mBuild(*items):
    """构建完整 <m:oMath> -> XML 字符串"""
    omath = etree.Element(_m('oMath'))
    for item in items:
        if isinstance(item, str):
            omath.append(_mr(item))
        elif isinstance(item, list):
            for sub in item:
                omath.append(sub)
        else:
            omath.append(item)
    return etree.tostring(omath, encoding='unicode')

def add_equation(num, omml_str):
    """插入 Word 原生公式 + 右对齐编号"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(7.32), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.64), WD_TAB_ALIGNMENT.RIGHT)

    # 居中公式
    run_eq = p.add_run('\t')
    omml_elem = etree.fromstring(omml_str.encode('utf-8'))
    run_eq._element.append(omml_elem)

    # 右对齐编号
    run_num = p.add_run(f'\t({num})')
    run_num.font.size = Pt(10.5)
    run_num.font.name = 'Times New Roman'
    return p

doc = Document()

# ============================================================
# 页面设置 (A4, 标准边距)
# ============================================================
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.18)
section.right_margin  = Cm(3.18)

# ============================================================
# 样式定义 (匹配参考论文)
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符

# Heading 1 样式
h1_style = doc.styles['Heading 1']
h1_style.font.name = '黑体'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1_style.paragraph_format.space_before = Pt(6)
h1_style.paragraph_format.space_after = Pt(3)
h1_style.paragraph_format.first_line_indent = Cm(0)

# Heading 2 样式
h2_style = doc.styles['Heading 2']
h2_style.font.name = '黑体'
h2_style.font.size = Pt(10.5)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_style.paragraph_format.space_before = Pt(3)
h2_style.paragraph_format.space_after = Pt(2)
h2_style.paragraph_format.first_line_indent = Cm(0)

# Heading 3 样式
h3_style = doc.styles['Heading 3']
h3_style.font.name = '黑体'
h3_style.font.size = Pt(10.5)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0, 0, 0)
h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h3_style.paragraph_format.space_before = Pt(2)
h3_style.paragraph_format.space_after = Pt(1)
h3_style.paragraph_format.first_line_indent = Cm(0)

# Title 样式
title_style = doc.styles['Title']
title_style.font.name = '黑体'
title_style.font.size = Pt(22)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)
title_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(4)
title_style.paragraph_format.first_line_indent = Cm(0)

# ============================================================
# 辅助函数
# ============================================================
def add_paragraph(text, style_name='Normal', bold=False, font_name=None,
                  font_size=None, alignment=None, first_line_indent=True,
                  space_after=None):
    """添加段落并设置格式"""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading(text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    return h

def add_chinese_title(text):
    """添加中文论文标题 (22pt 黑体加粗居中)"""
    p = doc.add_paragraph(style='Title')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_english_title(text):
    """添加英文论文标题 (~14.5pt Times New Roman 加粗居中)"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14.5)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def add_author_line(text, font_size=10.5):
    """添加作者行 (居中)"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_affiliation(text):
    """添加单位行"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_abstract_cn(text):
    """中文摘要"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run('摘要：')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = '黑体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = '宋体'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_keywords_cn(text):
    """中文关键词"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run('关键词：')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = '黑体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = '宋体'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_clc_line():
    """中图分类号行"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('中图分类号：  TM912\t文献标志码：A\t        文章编号：')
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_abstract_en(text):
    """英文摘要"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run('Abstract: ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Times New Roman'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def add_keywords_en(text):
    """英文关键词"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run('Keywords: ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Times New Roman'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def _inline_omml_str(omml_elem):
    """Convert an OMML element to XML string for inline insertion."""
    omath = etree.Element(_m('oMath'))
    if isinstance(omml_elem, str):
        omath.append(_mr(omml_elem))
    else:
        omath.append(omml_elem)
    return etree.tostring(omath, encoding='unicode')

def _build_inline_sub(base, sub):
    """Build inline OMML string for a subscript symbol."""
    base_map = {'V': 'V', 'R': 'R', 'K': 'K', 'Q': 'Q', 'C': 'C',
                'tau': 'τ', 'G': 'G', 'sigma': 'σ'}
    sub_map = {'t': 't', '0': '0', '1': '1', 'sd': 'sd', 'n': 'n', 'v': 'v'}
    b = base_map.get(base, base)
    s = sub_map.get(sub, sub)
    return _inline_omml_str(_mSub(b, s))

# Cache of pre-built inline OMML strings
_INL_CACHE = {}
def _inl_sub(base, sub):
    """Cached inline subscript OMML string."""
    key = f'{base}_{sub}'
    if key not in _INL_CACHE:
        _INL_CACHE[key] = _build_inline_sub(base, sub)
    return _INL_CACHE[key]

def _inl_text(txt):
    """Cached inline plain math text OMML string."""
    key = f'__text__{txt}'
    if key not in _INL_CACHE:
        _INL_CACHE[key] = _inline_omml_str(_mr(txt))
    return _INL_CACHE[key]

import re

# Whitelist of valid subscript symbols for inline math
_VALID_SUB_SYMBOLS = {
    # (base, sub) pairs — both ASCII and Greek letter variants
    ('V', 't'), ('V', '1'), ('V', 'sd'),
    ('R', '0'), ('R', '1'),
    ('C', '1'),
    ('K', 'sd'), ('Q', 'n'),
    ('tau', '1'), ('tau', 'sd'),  # ASCII tau
    ('τ', '1'), ('τ', 'sd'),     # Greek tau
    ('sigma', 'v'), ('σ', 'v'),  # ASCII/Greek sigma
    ('G', 'sd'),
}

def _auto_math_replacer(m):
    """Given a regex match, return [('sub', base, sub), ...] or [('math', text), ...]."""
    full = m.group(0)
    # G_sd(s) — special case
    if full == 'G_sd(s)':
        return [('sub', 'G', 'sd'), ('text', '(s)')]
    # Underscore notation: V_t, R_0, tau_sd, etc.
    m2 = re.match(r'^([A-Za-zτσΔ]+)_([a-z0-9]+)$', full)
    if m2:
        base, sub = m2.group(1), m2.group(2)
        if (base, sub) in _VALID_SUB_SYMBOLS:
            return [('sub', base, sub)]
        return None
    # Unicode subscript: V₁, R₀, R₁, C₁, τ₁
    uni_map = {'₁': '1', '₂': '2', '₃': '3', '₀': '0', '₆': '6', '₇': '7', '₈': '8'}
    m3 = re.match(r'^([A-Za-zτ])([₁₂₃₀₆₇₈]+)$', full)
    if m3:
        base = m3.group(1)
        sub = ''.join(uni_map.get(c, c) for c in m3.group(2))
        if (base, sub) in _VALID_SUB_SYMBOLS:
            return [('sub', base, sub)]
        return None
    # Δt
    if full == 'Δt':
        return [('math', 'Δt')]
    # OCV, SOC, UKF, BMS
    if full in ('OCV', 'SOC'):
        return [('math', full)]
    return None

# Compiled regex: whitelist-based patterns
_MATH_SYMBOLS_RE = re.compile(
    r'G_sd\(s\)|'  # special case
    r'(?:V|R|C|K|Q|tau|τ|sigma|σ|G)_(?:t|[01]|sd|n|v)|'  # known underscore symbols
    r'[VRC][₀₁]|τ[₁₂]|'  # known Unicode subscript symbols
    r'Δt|'  # Delta t
    r'(?<![A-Za-z])OCV(?![A-Za-z])|(?<![A-Za-z])SOC(?![A-Za-z])'  # acronyms
)

def add_body(text):
    """正文段落, 自动检测内联数学符号并转换为 Word 原生公式."""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.74)

    # Split text by math symbol matches
    parts = []
    last_end = 0
    for m in _MATH_SYMBOLS_RE.finditer(text):
        # Text before this match
        if m.start() > last_end:
            parts.append(('text', text[last_end:m.start()]))
        # The math symbol
        result = _auto_math_replacer(m)
        if result:
            parts.extend(result)
        else:
            parts.append(('text', m.group(0)))
        last_end = m.end()
    # Remaining text
    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    # Build paragraph with mixed text/OMML runs
    for item in parts:
        part_type = item[0]
        if part_type == 'text':
            run = p.add_run(item[1])
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif part_type == 'sub':
            run = p.add_run()
            omml_str = _inl_sub(item[1], item[2])
            omml_elem = etree.fromstring(omml_str.encode('utf-8'))
            run._element.append(omml_elem)
        elif part_type == 'math':
            run = p.add_run()
            omml_str = _inl_text(item[1])
            omml_elem = etree.fromstring(omml_str.encode('utf-8'))
            run._element.append(omml_elem)
    return p

def add_figure(fig_num, cn_text, en_text, filename, width_inches=5.5):
    """插入图片+中英文图题"""
    img_path = os.path.join(os.path.dirname(__file__), filename)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph(style='Normal')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.first_line_indent = Cm(0)
        p_img.paragraph_format.space_before = Pt(6)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_inches))
    # 中文图题
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run(f'图 {fig_num}  {cn_text}')
    r1.font.size = Pt(9)
    r1.font.name = '宋体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 英文图题
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run(f'Fig. {fig_num}  {en_text}')
    r2.font.size = Pt(9)
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2.bold = True
    return p2


def add_figure_caption(fig_num, cn_text, en_text):
    """图题 (中文 + 英文) — 仅文字，不插入图片"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run(f'图 {fig_num}  {cn_text}')
    r1.font.size = Pt(9)
    r1.font.name = '宋体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run(f'Fig. {fig_num}  {en_text}')
    r2.font.size = Pt(9)
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2.bold = True
    return p2

def add_table_caption(tab_num, cn_text, en_text):
    """表题 (中文 + 英文)"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f'表 {tab_num}  {cn_text}')
    r1.font.size = Pt(9)
    r1.font.name = '黑体'
    r1.bold = True
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p2 = doc.add_paragraph(style='Normal')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run(f'Table {tab_num}  {en_text}')
    r2.font.size = Pt(9)
    r2.font.name = 'Times New Roman'
    r2.bold = True
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p2

def add_table_with_data(headers, rows, col_widths=None):
    """添加表格并填充数据"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def _setup_ref_numbering():
    """在文档中添加 [1] 格式的编号定义。"""
    numbering_part = doc.part.numbering_part
    nsmap = numbering_part._element.nsmap

    # 找到所有已用的 abstractNumId 和 numId，取最大值 + 1
    max_id = 0
    for el in numbering_part._element:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag in ('abstractNum', 'num'):
            aid = int(el.get(qn('w:abstractNumId'), el.get(qn('w:numId'), '0')))
            if aid > max_id:
                max_id = aid
    abs_id = max_id + 1
    num_id = abs_id

    # abstractNum
    abn = OxmlElement('w:abstractNum')
    abn.set(qn('w:abstractNumId'), str(abs_id))

    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')

    start_el = OxmlElement('w:start')
    start_el.set(qn('w:val'), '1')
    lvl.append(start_el)

    nf = OxmlElement('w:numFmt')
    nf.set(qn('w:val'), 'decimal')
    lvl.append(nf)

    lt = OxmlElement('w:lvlText')
    lt.set(qn('w:val'), '[%1]')
    lvl.append(lt)

    lj = OxmlElement('w:lvlJc')
    lj.set(qn('w:val'), 'left')
    lvl.append(lj)

    pp = OxmlElement('w:pPr')
    ind_el = OxmlElement('w:ind')
    ind_el.set(qn('w:left'), '440')
    ind_el.set(qn('w:hanging'), '440')
    pp.append(ind_el)
    lvl.append(pp)

    abn.append(lvl)
    numbering_part._element.append(abn)

    # num 实例
    num_el = OxmlElement('w:num')
    num_el.set(qn('w:numId'), str(num_id))
    aid_ref = OxmlElement('w:abstractNumId')
    aid_ref.set(qn('w:val'), str(abs_id))
    num_el.append(aid_ref)
    numbering_part._element.append(num_el)

    return num_id


_ref_numId = None  # 在文档构建阶段赋值


def add_ref(text):
    """参考文献条目 — 10.5pt, 单倍行距, [N] 自动编号"""
    global _ref_numId
    if _ref_numId is None:
        _ref_numId = _setup_ref_numbering()

    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.line_spacing = 1.0

    # 挂接编号
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), '0')
    numPr.append(ilvl_el)
    nid_el = OxmlElement('w:numId')
    nid_el.set(qn('w:val'), str(_ref_numId))
    numPr.append(nid_el)
    pPr.append(numPr)

    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


# ============================================================
# 正文内容
# ============================================================

# —————— 中文标题 ——————
add_chinese_title('基于固相扩散慢弛豫修正的锂离子电池')
add_chinese_title('SOC高精度估计方法')

# 作者
add_author_line('侯宇欣¹，XXX²，XXX¹*')

# 单位
add_affiliation('（1. XXX大学 XXX学院，成都 610000；2. XXX大学 XXX学院，XX 000000）')

# 中文摘要
add_abstract_cn(
    '针对传统一阶RC等效电路模型忽略固相扩散慢弛豫特性导致锂离子电池荷电状态(SOC)'
    '估算精度不足的问题，提出一种轻量化模型改进方案。在不增加RC支路、不改变模型阶数的前提下，'
    '引入一阶惯性环节作为固相扩散修正项，仅需整定2个额外参数(τ_sd、K_sd)，实现电池快极化'
    '(电化学极化)与慢扩散(固相浓差极化)特性的解耦表征。在此基础上，结合无迹卡尔曼滤波(UKF)'
    '算法融合安时积分法与端电压观测值，建立改进模型的状态方程与观测方程。仿真结果表明，与'
    '传统一阶RC+UKF方案相比，改进模型的最大SOC估算误差从1.86%降至1.36%，平均误差从1.09%'
    '降至0.48%，均方根误差(RMSE)从1.23%降至0.60%，三项精度指标均满足预设目标(最大误差≤1.5%、'
    '平均误差≤0.8%、RMSE≤1.0%)。进一步提出基于脉冲-静置实验的参数离线辨识方法，τ_sd与K_sd'
    '的辨识误差分别为0.7%和0.2%，解决了新增参数的工程获取问题。该方法在保留一阶RC模型结构'
    '简单、计算量小的核心优势的同时，显著提升了SOC估算精度，为车载BMS等实时性要求较高的场景'
    '提供了简单可行的解决方案。'
)

# 关键词
add_keywords_cn('锂离子电池；荷电状态(SOC)；固相扩散；无迹卡尔曼滤波；等效电路模型；参数辨识')

# 中图分类号
add_clc_line()

# —————— 英文标题 ——————
add_english_title('High-Accuracy SOC Estimation for Lithium-Ion Batteries')
add_english_title('Based on Solid-Phase Diffusion Slow-Relaxation Compensation')

# 英文作者
add_author_line('HOU Yuxin¹, XXX², XXX¹*', font_size=10)

# 英文单位
add_affiliation('(1. School of XXX, XXX University, Chengdu 610000, China; '
                '2. School of XXX, XXX University, XX 000000, China)')

# 英文摘要
add_abstract_en(
    'To address the problem of insufficient State of Charge (SOC) estimation accuracy in '
    'traditional first-order RC equivalent circuit models due to the neglect of solid-phase '
    'diffusion slow-relaxation characteristics, this paper proposes a lightweight model '
    'improvement scheme. Without adding RC branches or increasing model order, a first-order '
    'inertia link is introduced as a solid-phase diffusion correction term, requiring only '
    'two additional tunable parameters (τ_sd, K_sd) to achieve decoupled characterization of '
    'fast polarization (electrochemical polarization) and slow diffusion (solid-phase '
    'concentration polarization). Based on the improved model, an Unscented Kalman Filter (UKF) '
    'algorithm fuses ampere-hour integration with terminal voltage observations. Simulation '
    'results show that compared with the traditional first-order RC + UKF scheme, the improved '
    'model reduces the maximum SOC estimation error from 1.86% to 1.36%, the mean error from '
    '1.09% to 0.48%, and the RMSE from 1.23% to 0.60%, with all three accuracy metrics meeting '
    'the preset targets (max error ≤1.5%, mean error ≤0.8%, RMSE ≤1.0%). Furthermore, an offline '
    'parameter identification method based on pulse-relaxation experiments is proposed, achieving '
    'identification errors of 0.7% for τ_sd and 0.2% for K_sd, thereby solving the engineering '
    'problem of obtaining the additional parameters. The proposed method retains the core advantages '
    'of the first-order RC model—simple structure and low computational cost—while significantly '
    'improving SOC estimation accuracy, providing a practical solution for real-time applications '
    'such as onboard Battery Management Systems (BMS).'
)

# 英文关键词
add_keywords_en('lithium-ion battery; State of Charge (SOC); solid-phase diffusion; '
                'Unscented Kalman Filter; equivalent circuit model; parameter identification')


# ============================================================
# 1 引言
# ============================================================
add_heading('1 引言', level=1)

add_body(
    '锂离子电池因高能量密度和长循环寿命等优势，已广泛应用于电动汽车和储能系统[1-2]。'
    '荷电状态(SOC)的精确估计是电池管理系统(BMS)最关键的功能之一，直接影响续航预测和充放电'
    '策略[3-4]。然而SOC无法直接测量，须通过电压、电流等间接估计，是该领域的研究难点[5]。'
)

add_body(
    'SOC估计方法可分为安时积分法、基于模型的方法和数据驱动方法三类[6]。安时积分法简单但'
    '对累积误差敏感[7]；数据驱动方法泛化能力有限[8]；基于等效电路模型(ECM)结合滤波算法的'
    '闭环估计方法兼顾精度与鲁棒性，是工程应用的主流路线[9]。一阶RC(Thevenin)模型因参数少、'
    '实时性好，在车载BMS中应用最广[10-11]。然而，传统一阶RC模型用单一RC支路同时表征电化学极化'
    '(快动态，τ≈10~100s)和固相扩散浓差极化(慢动态，τ≈100~1000s)，在静置阶段无法准确描述'
    '电压慢弛豫，导致SOC估计出现系统性偏差。'
)

add_body(
    '现有改进方案沿两条路线展开：增加RC阶数[12]或采用电化学/分数阶模型[13-14]，但均面临'
    '参数增多、计算负担加重的问题，难以满足嵌入式BMS实时性约束。本文提出轻量化改进方案：'
    '在不增加RC支路的前提下，引入一阶惯性环节G_sd(s)=K_sd/(τ_sd·s+1)作为固相扩散修正项，'
    '仅增加2个可辨识参数，实现快慢极化解耦表征。主要贡献包括：(1)在不增加模型阶数的条件下'
    '实现固相扩散慢弛豫的精确表征；(2)提出脉冲-静置两阶段参数离线辨识方法，τ_sd和K_sd辨识'
    '误差分别低至0.7%和0.2%；(3)改进模型SOC估计RMSE从1.23%降至0.60%，三项精度指标全部达标。'
)

# ============================================================
# 2 相关工作
# ============================================================
add_heading('2 相关工作', level=1)

add_heading('2.1 锂离子电池内部结构与固相扩散机理', level=2)

add_body(
    '锂离子电池内部结构及工作原理如图1所示。电池由正极(NMC)、负极(石墨)、隔膜和电解液组成。'
    '放电时，Li⁺从负极石墨层间脱嵌，经电解液穿过隔膜嵌入正极NMC晶格，电子则通过外电路由负极'
    '流向正极，对外做功；充电过程反向。正负极活性材料均为微米级多孔颗粒，Li⁺在颗粒内部的固相'
    '扩散系数(D_s≈10⁻¹¹~10⁻¹³ cm²/s)远小于液相扩散系数(D_e≈10⁻⁵~10⁻⁶ cm²/s)，是电池动力学'
    '的限速步骤。固相扩散弛豫时间(τ_sd≈100~1000s)远超电化学极化时间常数(τ₁≈10~100s)，构成'
    '"固相扩散慢弛豫"的物理来源[15-16]。该慢过程的准确表征是SOC估计精度的关键挑战。'
)

add_figure(1, '锂离子电池内部结构与离子/电子流向示意图',
           'Schematic of Li-ion battery internal structure and ion/electron transport pathways, '
           'highlighting solid-phase diffusion within electrode particles',
           'fig_battery_structure.png', width_inches=5.0)

add_heading('2.2 锂离子电池极化特性与等效电路建模', level=2)

add_body(
    '锂离子电池的极化行为包括欧姆极化(R₀表征)、电化学极化(τ≈10~100s)和固相浓差极化'
    '(τ≈100~1000s)三类[15-16]。传统一阶RC(Thevenin)模型仅用一个RC支路同时拟合快慢两类极化，'
    '方程如下[17]：'
)

# exp(-Δt/τ_1) — shared between Eq(2) and Eq(7)
_exp_neg_tau1 = _mFunc('exp', '-',
    _mFrac(_mr('Δt'), _mSub('τ', '1')))

add_equation(1, _mBuild(
    _mSub('V', 't'),
    ' = ', 'OCV', _mParen('SOC'),
    ' - ', _mSub('V', '1'),
    ' - ', 'I', ' · ', _mSub('R', '0'),
))
add_equation(2, _mBuild(
    _mSub('V', '1'), _mParen('k+1'),
    ' = ',
    _mSub('V', '1'), _mParen('k'),
    ' · ', _exp_neg_tau1,
    ' + ', 'I', _mParen('k'),
    ' · ', _mSub('R', '1'),
    ' · ', _mParen('1', ' - ', _exp_neg_tau1),
))

add_body(
    '其中V_t为端电压，OCV为开路电压，V₁为极化电压，R₀为欧姆内阻，τ₁=R₁C₁为RC时间常数。'
    '由于单一RC支路只能反映一种时间尺度，τ₁的辨识值只能近似综合两种极化的响应，无法解耦快慢'
    '动力学差异，导致静置阶段端电压预测偏差，进而使SOC估计产生漂移[18-19]。'
)

add_heading('2.3 基于UKF的SOC估计方法', level=2)

add_body(
    '无迹卡尔曼滤波(UKF)通过无迹变换生成2n+1个确定性Sigma点直接传播状态分布，避免扩展卡尔曼'
    '滤波(EKF)的线性化误差，对电池OCV-SOC强非线性系统估计精度更高[20-22]。UKF核心步骤包括：'
    'Sigma点生成→状态预测→观测预测→卡尔曼增益计算→状态更新[23]。当模型存在结构缺陷时，'
    'UKF虽能通过电压反馈部分补偿，但模型遗漏的物理效应最终会映射为SOC系统偏差[24-25]。'
)

# ============================================================
# 3 方法
# ============================================================
add_heading('3 方法', level=1)

add_body(
    '本文方法整体框架如图2所示，主要包括三个层次：物理系统层(电池+测量)、模型层(改进ECM+参数辨识)'
    '和估计算法层(UKF+新息反馈)，各模块的详细设计如下。'
)

add_figure(2, '方法论框架：改进等效电路模型与UKF联合SOC估计流程',
           'Methodology overview: Improved ECM + UKF framework for SOC estimation',
           'fig_methodology.png', width_inches=5.8)

add_heading('3.1 改进一阶RC模型设计', level=2)

add_body(
    '本文在原有一阶RC模型基础上引入一阶惯性环节G_sd(s)=K_sd/(τ_sd·s+1)作为固相扩散修正项，'
    '将传统"单RC同时表征快慢极化"升级为"RC(快极化)+惯性环节(慢扩散)"的解耦结构，不增加'
    'RC支路，仅新增2个参数。'
)

# exp(-Δt/τ_sd) — for Eq(4)
_exp_neg_tsd = _mFunc('exp', '-',
    _mFrac(_mr('Δt'), _mSub('τ', 'sd')))

add_equation(4, _mBuild(
    _mSub('V', 'sd'), _mParen('k+1'),
    ' = ',
    _mSub('V', 'sd'), _mParen('k'),
    ' · ', _exp_neg_tsd,
    ' + ', 'I', _mParen('k'),
    ' · ', _mSub('K', 'sd'),
    ' · ', _mParen('1', ' - ', _exp_neg_tsd),
))

add_body(
    '改进模型状态向量x=[SOC, V₁, V_sd]ᵀ，端电压方程：'
)

add_equation(5, _mBuild(
    _mSub('V', 't'),
    ' = ', 'OCV', _mParen('SOC'),
    ' - ', _mSub('V', '1'),
    ' - ', _mSub('V', 'sd'),
    ' - ', 'I', ' · ', _mSub('R', '0'),
))

add_body(
    '其中V₁(τ₁≈37.5s)表征电化学极化快动态，V_sd(τ_sd≈280s)独立表征固相扩散慢弛豫，'
    '实现了两种动力学的解耦。'
)

add_heading('3.2 UKF状态估计', level=2)

add_body(
    '状态方程由安时积分耦合双极化动力学构成，观测方程为式(5)：'
)

add_equation(6, _mBuild(
    'SOC', _mParen('k+1'),
    ' = ', 'SOC', _mParen('k'),
    ' - ', 'Δt', ' · ', 'I', _mParen('k'), ' / ', _mSub('Q', 'n'),
))
add_equation(7, _mBuild(
    _mSub('V', '1'), _mParen('k+1'),
    ' = ',
    _mSub('V', '1'), _mParen('k'),
    ' · ', _exp_neg_tau1,
    ' + ', 'I', _mParen('k'),
    ' · ', _mSub('R', '1'),
    ' · ', _mParen('1', ' - ', _exp_neg_tau1),
))

add_body(
    'OCV-SOC采用三阶多项式：OCV(SOC)=0.9SOC³−1.5SOC²+1.8SOC+3.0 (3.0~4.2V)，'
    '特性曲线如图3所示。'
    'UKF参数：α=1.0, β=2.0, κ=0；Q=diag(1×10⁻⁸, 1×10⁻⁶, 5×10⁻⁷)；R=4×10⁻⁶(σ_v=2mV)；'
    'Δt=0.1s。'
)

add_figure(3, 'OCV-SOC特性曲线（NMC/石墨体系）',
           'OCV-SOC characteristic curve (NMC/Graphite)',
           'fig_ocv_soc.png')

add_heading('3.3 参数离线辨识方法', level=2)

add_body(
    '对新增参数τ_sd和K_sd，提出脉冲-静置两阶段辨识方法。在已知SOC下施加1C恒流脉冲600s后'
    '静置1800s，记录端电压弛豫。静置阶段I=0时V_t(t)=OCV−V₁(0)e^(−t/τ₁)−V_sd(0)e^(−t/τ_sd)。'
    '由于τ_sd>>τ₁，约200s后V₁基本归零，端电压变化仅由V_sd衰减主导。'
)

add_body(
    '阶段一：利用已知τ₁、R₁扣除V₁贡献，得V_corrected=OCV−V_t−V₁(t)；阶段二：对V_corrected'
    '做单指数非线性最小二乘拟合得τ_sd和V_sd(0)，反算K_sd=V_sd(0)/[I·(1−e^(−t_pulse/τ_sd))]。'
    '方法仅需常规充放电设备，单次实验即可同时获取两个参数。'
)

# ============================================================
# 4 实验设计
# ============================================================
add_heading('4 实验设计', level=1)

add_body(
    '基于Python搭建仿真平台，采用NMC/石墨体系电池典型参数，如表1所示。OCV-SOC采用三阶多项式：'
    'OCV(SOC)=0.9SOC³−1.5SOC²+1.8SOC+3.0 (3.0~4.2V)，UKF初始SOC=80%，随机种子固定(42)。'
)

# 表1 — 电池模型参数
add_table_caption(1, '电池模型仿真参数', 'Battery Model Simulation Parameters')
add_table_with_data(
    ['参数', '符号', '数值', '单位'],
    [
        ['额定容量', 'Q_n', '2.5', 'Ah'],
        ['欧姆内阻', 'R₀', '0.025', 'Ω'],
        ['极化电阻', 'R₁', '0.015', 'Ω'],
        ['极化电容', 'C₁', '2500', 'F'],
        ['RC时间常数', 'τ₁=R₁C₁', '37.5', 's'],
        ['固相扩散时间常数', 'τ_sd', '280', 's'],
        ['固相扩散增益', 'K_sd', '0.008', '—'],
        ['采样周期', 'Δt', '0.1', 's'],
        ['电压测量噪声标准差', 'σ_v', '2.0', 'mV'],
    ]
)
doc.add_paragraph()

add_body(
    '设置两种工况：(1)工况一——1C放电1500s→静置900s→1C充电1500s→静置900s，覆盖充放电快响应'
    '与静置慢弛豫，评估SOC估计性能；(2)工况二——1C放电脉冲600s→静置1800s，用于τ_sd和K_sd离线'
    '辨识。三组对比模型：(1)传统一阶RC+UKF(n=2，基线)；(2)改进模型+UKF(n=3，本文主推方案)；'
    '(3)自适应模型+UKF(n=4，τ_sd在线估计，初值300s)。评估指标：最大误差(MaxE≤1.5%)、'
    '平均误差(MAE≤0.8%)、RMSE(≤1.0%)，以及端电压RMSE(mV)作为辅助指标。'
)

# ============================================================
# 5 结果与分析
# ============================================================
add_heading('5 结果与分析', level=1)

add_body(
    '本节遵循"整体精度对比→关键现象分析→参数辨识验证→自适应扩展评估"的递进逻辑展开。'
    '首先在标准工况下对比三组模型的SOC估计性能；随后针对静置阶段的SOC漂移问题进行深入分析；'
    '然后验证参数离线辨识方法的精度；最后评估自适应UKF框架的可行性。'
)

add_heading('5.1 三模型SOC估计精度对比', level=2)

add_body(
    '如表2所示，三组模型在工况一下的SOC估计精度对比结果汇总如下。传统模型的三项精度指标全部未达标'
    '(0/3 PASS)，其中最大误差1.86%、平均误差1.09%、RMSE 1.23%，证实了模型结构缺陷(忽略'
    '固相扩散)导致的系统性精度损失。改进模型在三项精度指标上均达到预设目标(ALL PASS)，'
    '最大误差降至1.36%(降幅26.7%)，平均误差降至0.48%(降幅56.6%)，RMSE降至0.60%(降幅50.9%)，'
    '充分验证了固相扩散修正项的有效性。自适应模型的最大误差为1.86%(与传统模型持平，未达标)，'
    '但平均误差(0.55%)和RMSE(0.73%)均达标，且τ_sd在线估计结果(收敛至293.1s)与真值(280s)'
    '的偏差仅4.7%，展示了自适应框架的可行性与当前局限。'
    '三模型SOC估计曲线及误差的详细对比如图4所示，误差分布统计特性见图5。'
)

# 表2 — SOC估计误差对比
add_table_caption(2, '三模型SOC估计精度对比(工况一)', 'Three-Model SOC Estimation Accuracy Comparison (Profile I)')
add_table_with_data(
    ['模型', '最大误差(%)', '平均误差(%)', 'RMSE(%)', 'Vt RMSE(mV)', '达标情况'],
    [
        ['传统模型', '1.8603', '1.0946', '1.2283', '0.9829', '0/3 FAIL'],
        ['改进模型', '1.3629', '0.4750', '0.6027', '1.0819', 'ALL PASS'],
        ['自适应模型', '1.8603', '0.5506', '0.7289', '1.1547', '2/3 PASS'],
        ['目标值', '≤1.5', '≤0.8', '≤1.0', '—', '—'],
    ]
)
doc.add_paragraph()

add_figure(4, '三模型SOC估计综合对比 (a)电流 (b)端电压 (c)SOC估计 (d)SOC误差 (e)V_sd跟踪 (f)τ_sd在线自适应',
           'Three-model SOC estimation comprehensive comparison (a)Current (b)Terminal voltage '
           '(c)SOC estimation (d)SOC error (e)V_sd tracking (f)Online tau_sd adaptation',
           'three_way_comparison.png')

add_figure(5, '三模型SOC估计误差分布对比（直方图+高斯拟合）',
           'SOC estimation error distribution comparison (histogram + Gaussian fit)',
           'fig_error_distribution.png')

add_heading('5.2 端电压估计与SOC跟踪', level=2)

add_body(
    '三组模型的端电压估计RMSE均处于较低水平(0.98~1.15mV)，表明UKF能够有效融合电压测量新息。'
    '然而相近的端电压精度并不等同SOC精度：传统模型RMSE(1.23%)约为改进模型(0.60%)的两倍。'
    '根本原因在于传统模型缺少V_sd修正项，UKF将静置阶段固相扩散引起的电压弛豫错误映射为SOC修正，'
    '而改进模型正确分配了电压残差的物理来源。端电压各分量贡献的分解如图6所示。'
)

add_figure(6, '端电压分量分解（OCV、V1、V_sd、I·R0各自贡献）',
           'Terminal voltage component decomposition (OCV, V1, V_sd, I*R0 contributions)',
           'fig_voltage_decomposition.png')

add_heading('5.3 静置阶段SOC漂移抑制分析', level=2)

add_body(
    '静置阶段是模型差异最显著的观察窗口。放电结束后的900s静置期间，传统模型因缺少V_sd衰减机制，'
    '将约13~15mV的电压回升错误解释为SOC变化，导致估计偏离1.5%~2%；改进模型通过V_sd独立追踪'
    '慢弛豫，SOC误差控制在0.3%以内。V_sd幅值(约20mV)虽远小于V₁(37.5mV)和I·R₀(62.5mV)，'
    '但因其τ_sd(280s)远超τ₁(37.5s)，在静置阶段累积效应显著——这正是"慢弛豫"的物理本质。'
    '静置阶段SOC漂移抑制对比如图7所示。'
)

add_figure(7, '静置阶段SOC漂移抑制效果对比',
           'Comparison of SOC drift suppression during rest periods',
           'rest_period_detail.png')

add_heading('5.4 参数离线辨识结果', level=2)

add_body(
    '参数离线辨识结果如表3所示，辨识过程的可视化见图8。τ_sd辨识值282.1s(真值280s, 误差0.7%)，K_sd辨识值0.0080'
    '(真值0.0080, 误差0.2%)，辨识质量为GOOD。精度得益于：600s长脉冲使V_sd充分建立(约17.7mV)、'
    '0.5mV RMS低噪声、以及两阶段非线性拟合策略对噪声的鲁棒性优于传统对数-线性回归。'
    '该方法仅需常规充放电设备，单次实验即可获取两个参数，为BMS部署提供了可行途径。'
)

# 表3 — 参数辨识结果
add_table_caption(3, '脉冲-静置法参数辨识结果', 'Parameter Identification Results via Pulse-Relaxation Method')
add_table_with_data(
    ['参数', '真值', '辨识值', '绝对误差', '相对误差(%)'],
    [
        ['τ_sd (s)', '280.0', '282.1', '2.1', '0.7'],
        ['K_sd', '0.0080', '0.0080', '0.000016', '0.2'],
        ['V_sd0 (mV)', '17.65', '17.58', '0.07', '—'],
    ]
)
doc.add_paragraph()

add_figure(8, '脉冲-静置法参数辨识过程 (a)脉冲电流 (b)端电压响应 (c)两阶段拟合 (d)参数辨识精度',
           'Parameter identification via pulse-relaxation method (a)Pulse current '
           '(b)Terminal voltage response (c)Two-stage fit (d)Parameter accuracy',
           'param_identification.png')

add_heading('5.5 自适应UKF性能分析', level=2)

add_body(
    '自适应模型(n=4, τ_sd在线估计)实现了平均误差0.55%和RMSE 0.73%(均达标)，τ_sd收敛至293.1s'
    '(与真值280s偏差4.7%)，验证了增广状态UKF框架的可行性。但其最大误差1.86%与传统模型持平，'
    '表明联合状态-参数估计增加了滤波器自由度与不确定性，暂态稳定性不及固定参数方案。两者互补：'
    '参数稳定时改进模型最优，参数漂移时自适应框架更具长期鲁棒性。'
)

# ============================================================
# 6 讨论与结论
# ============================================================
add_heading('6 讨论与结论', level=1)

add_body(
    '本文提出了一种基于固相扩散慢弛豫修正的锂离子电池SOC轻量化估计方法。仿真结果表明，改进模型'
    '较传统一阶RC+UKF方案，RMSE从1.23%降至0.60%(降幅50.9%)，最大误差从1.86%降至1.36%(降幅26.7%)，'
    '平均误差从1.09%降至0.48%(降幅56.6%)，三项精度指标全部达标。改进模型仅增加1个状态变量和'
    '2个可辨识参数，计算增量几乎为零，完整保留了传统一阶RC模型实时性好的工程优势。提出的脉冲-静置'
    '两阶段离线辨识方法，τ_sd和K_sd辨识误差分别低至0.7%和0.2%，仅需常规充放电设备即可完成。'
)

add_body(
    '本文仍存在以下局限：(1)仅基于仿真验证，需实物电池实验确认；(2)OCV-SOC采用固定多项式，未考虑'
    '温度和老化影响；(3)未在动态驾驶工况(UDDS、FUDS)下验证。未来工作将沿实物实验验证、温度/老化'
    '补偿机制、SOC-SOH联合估计和动态工况泛化方向展开。'
)

# ============================================================
# 参考文献
# ============================================================
add_heading('参考文献', level=2)

references = [
    'Wang S, Fernandez C, Yu C, et al. A review of battery state of charge estimation '
    'based on equivalent circuit models[J]. Journal of Energy Storage, 2025, 110: 115346.',

    'Lipu M S H, Hannan M A, Hussain A, et al. Hybrid and combined states estimation '
    'approaches for lithium-ion battery management system: Advancement, challenges and future '
    'directions[J]. Journal of Energy Storage, 2024, 92: 112128.',

    'Barik S, Saravanan B. Recent developments and challenges in state-of-charge estimation '
    'techniques for electric vehicle batteries: A review[J]. Sustainable Energy Technologies and '
    'Assessments, 2024, 66: 103895.',

    'Chen X, Liu Z, Wang Y, et al. Evolution of electrical vehicles, battery state '
    'estimation, and future research directions: A critical review[J]. IEEE Access, 2024, '
    '12: 158627-158646.',

    'Selvaraj V, Vairavasundaram I. A comprehensive review of state of charge estimation '
    'in lithium-ion batteries used in electric vehicles[J]. Journal of Energy Storage, 2023, '
    '72: 208777.',

    'Chen L, Wang Z, Lyu Z, et al. A comparative study of equivalent circuit models for '
    'lithium-ion batteries: Accuracy, complexity, and application perspectives[J]. IEEE '
    'Transactions on Transportation Electrification, 2022, 8(2): 2789-2800.',

    'Zhang Y, Huang Z, Liu X, et al. An improved adaptive cubature Kalman filter for state '
    'of charge estimation of lithium-ion batteries[J]. Journal of Power Sources, 2023, '
    '555: 232371.',

    'Wang Y, Chen Z, Zhang C. A comprehensive review on data-driven state-of-charge '
    'estimation of lithium-ion batteries: Methods, challenges, and future directions[J]. '
    'Journal of Power Sources, 2024, 612: 235252.',

    'Xiong R, Cao J, Yu Q, et al. Model-based state of charge estimation for lithium-ion '
    'batteries: Recent advances, challenges, and perspectives[J]. Renewable and Sustainable '
    'Energy Reviews, 2023, 183: 113472.',

    'Li W, Deng J, Pang H, et al. A comparative study of lithium-ion battery equivalent '
    'circuit models for state of charge estimation under dynamic conditions[J]. Energy, 2024, '
    '294: 130844.',

    'Yang J, Xia B, Shang Y, et al. Comparative study on equivalent circuit models for '
    'state of charge estimation of lithium-ion batteries under wide temperature range[J]. '
    'Electrochimica Acta, 2023, 449: 142184.',

    'Tian J, Xiong R, Yu Q. Fractional-order model-based state of charge estimation for '
    'lithium-ion batteries considering hysteresis and temperature effects[J]. IEEE Transactions '
    'on Power Electronics, 2023, 38(4): 5102-5113.',

    'Huang Z, Zhang Y, Liu X, et al. Online estimation of battery equivalent circuit model '
    'parameters and state of charge using decoupled least squares technique with moving window[J]. '
    'Journal of Energy Storage, 2024, 79: 110144.',

    'Guo D, Yang G, Feng X, et al. Physics-based fractional-order model with simplified '
    'parameter identification for lithium-ion batteries[J]. Journal of Energy Storage, 2023, '
    '72: 108596.',

    'Hu X, Feng F, Liu K, et al. State estimation for advanced battery management: Key '
    'challenges and future trends[J]. Renewable and Sustainable Energy Reviews, 2021, '
    '152: 111648.',

    'Rüther T, Hileman W, Plett G L, et al. Demystifying the distribution of relaxation '
    'times: A simulation-based investigation into the limits and possibilities of interpretation '
    'for lithium-ion batteries[J]. Journal of the Electrochemical Society, 2024, 171(6): 060508.',

    'Wei Z, Zhao J, Zou C, et al. Comparative study of methods for battery state of charge '
    'estimation using extended and unscented Kalman filters under uncertain initial conditions[J]. '
    'Energy, 2024, 288: 129769.',

    'Xu Z, Wang J, Yang L, et al. State of charge estimation of lithium-ion batteries based '
    'on adaptive square root unscented Kalman filter with full-parameter online identification[J]. '
    'Journal of Energy Storage, 2024, 77: 109926.',

    'Sun L, Liu G, Chen Z, et al. An adaptive state-of-charge estimation method for '
    'lithium-ion batteries using improved unscented Kalman filter under wide temperature range[J]. '
    'Applied Energy, 2023, 350: 121731.',

    'Khalid A, Kashif S A R, Ain N U, et al. Comparison of Kalman filters for state '
    'estimation based on computational complexity of Li-ion cells[J]. Energies, 2023, '
    '16(6): 2710.',

    'Lopetegi I, Plett G L, Trimboli M S, et al. A new battery SOC/SOH/eSOH estimation '
    'method using a PBM and interconnected SPKFs: Part I. SOC and internal variable estimation[J]. '
    'Journal of the Electrochemical Society, 2024, 171(3): 030519.',

    'Khodarahmi M, Maihami V. A review on Kalman filter models: From standard Kalman '
    'filter to unscented and multiple model Kalman filters[J]. Archives of Computational Methods '
    'in Engineering, 2023, 30(1): 727-747.',

    'Meng J, Ricco M, Luo G, et al. An overview and comparison of online implementable SOC '
    'estimation methods for lithium-ion battery[J]. IEEE Transactions on Industry Applications, '
    '2022, 58(3): 3953-3964.',

    'Shen P, Ouyang M, Lu L, et al. The co-estimation of state of charge and state of health '
    'for lithium-ion batteries based on improved adaptive unscented Kalman filter[J]. IEEE '
    'Transactions on Power Electronics, 2024, 39(2): 2479-2490.',

    'Ye M, Guo H, Xiong R, et al. Co-estimation of lithium-ion battery state-of-charge and '
    'state-of-health based on fractional-order modeling and dual adaptive unscented Kalman '
    'filter[J]. Journal of Energy Storage, 2023, 65: 107225.',
]

for ref in references:
    add_ref(ref)

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT_PATH)
print(f'论文已保存至: {OUTPUT_PATH}')
print(f'文件大小: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB')
